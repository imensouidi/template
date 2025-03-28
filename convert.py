import json
import re
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract
from openai import AzureOpenAI
from azure.keyvault.secrets import SecretClient
from azure.identity import DefaultAzureCredential
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.lib.utils import ImageReader
from flask import Flask, request, jsonify
import os
import logging
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from flask_cors import CORS
import tempfile
from dotenv import load_dotenv
from datetime import datetime, timedelta, timezone
from pdf2docx import Converter
from docx.shared import Inches
from docx.shared import Pt

# Nouveaux imports pour modifier le XML du DOCX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Charger les variables d'environnement depuis le fichier .env
load_dotenv()

# Configuration de l’application Flask
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": ["https://talent.heptasys.com", "http://localhost:4200"]}},
     allow_headers=["Content-Type", "Authorization", "X-Requested-With"])
logging.basicConfig(level=logging.INFO)

# Configuration d’Azure Key Vault
key_vault_name = 'AI-vault-hepta'
key_vault_uri = f"https://{key_vault_name}.vault.azure.net/"
credential = DefaultAzureCredential()
secret_client = SecretClient(vault_url=key_vault_uri, credential=credential)

api_key = secret_client.get_secret('AZUREopenaiAPIkey').value
azure_endpoint = secret_client.get_secret('AZUREopenaiENDPOINT').value
connect_string = secret_client.get_secret('connectstr').value

azure_openai_client = AzureOpenAI(
    api_key=api_key,
    api_version="2024-02-15-preview",
    azure_endpoint=azure_endpoint
)

# Configuration du Blob Storage
blob_service_client = BlobServiceClient.from_connection_string(connect_string)
container_name = "converted"

def get_account_key_from_connection_string(conn_str):
    parts = conn_str.split(';')
    for part in parts:
        if part.strip().startswith("AccountKey="):
            return part.split("=", 1)[1]
    return None

_account_key = None
if hasattr(blob_service_client.credential, "account_key"):
    _account_key = blob_service_client.credential.account_key
else:
    _account_key = get_account_key_from_connection_string(connect_string)

def extract_text(file_path):
    """
    Extrait le texte d'un fichier PDF, DOCX ou image (PNG/JPG).
    Utilise PyMuPDF pour PDF, python-docx pour DOCX, et pytesseract pour les images.
    """
    try:
        if file_path.lower().endswith(".pdf"):
            with fitz.open(file_path) as doc:
                return " ".join(page.get_text() for page in doc)
        elif file_path.lower().endswith(".docx"):
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            image = Image.open(file_path)
            return pytesseract.image_to_string(image)
    except Exception as e:
        logging.error(f"Erreur lors de l'extraction du texte : {e}")
        return None

def extract_info_to_json(text):
    """
    Appelle AzureOpenAI pour extraire les informations du CV.
    """
    # Exemple de JSON attendu :
    json_format = """
{
  "job_title": "",
  "full_name": "",
  "years_of_experience": "",
  "contact_information": {
    "phone": "",
    "email": "",
    "website": ""
  },
  "education": [
    {
      "degree": "",
      "institution": "",
      "year_of_completion": ""
    }
  ],
  "professional_experience": [
    {
      "company_name": "",
      "date_range": "",
      "mission": "",
      "tasks": [""],
      "tech_tools": [""]
    }
  ],
  "skills": [],
  "certifications": []
}
    """
    
    # Prompt renforcé pour extraire un titre et respecter l’ordre des expériences
    prompt = f"""
You are a helpful assistant that extracts specific information from a résumé (CV) written in French.

IMPORTANT:
- The text is in French. DO NOT translate it.
- If the résumé has a distinct title (e.g. "Titre du CV", "Management de la maintenance", etc.), 
  please store it in the "job_title" field in the JSON.
- Extract the technical skills information from the CV and return it in the "skills" field.
  If the skills are organized in categories (i.e., each category is labeled with a title followed by a colon and a list of skills),
  return a dictionary where the keys are the category names and the values are the skills (each skill separated by a comma).
  If the skills are not organized in categories, return them as a single comma-separated string.
- Extract the certifications from the CV and return them in the "certifications" field as a list of strings.
- Preserve the exact order of the professional experiences as they appear in the original CV. 
  The first experience in the CV should remain the first in the JSON output, the second remains the second, etc. 
  Do not reorder or regroup them.
- For each entry in "professional_experience", extract exactly the following fields as they appear in the CV:
  "date_range", "company_name", and "mission".
  IMPORTANT: For "mission", extract only the job title – i.e. the first line immediately following 
  the date range and company name – and do not include any additional descriptive text.
- Return only valid JSON (no extra text or symbols).

Extract the following information from the text (in French):
{text}

Format the result as JSON according to the example below:
{json_format}

Do not include any extra symbols.
    """

    try:
        logging.info("=== Prompt envoyé à l'API ===")
        logging.info(prompt)
        
        response = azure_openai_client.completions.create(
            model="IndexSelector",  # Remplacez si nécessaire par le nom de déploiement effectif (ex: "gpt-35-turbo")
            prompt=prompt,
            max_tokens=3000,
            temperature=0
        )
        
        raw_json_text = response.choices[0].text.strip()
        logging.info("=== Réponse brute de l'API AzureOpenAI ===")
        logging.info(raw_json_text)
        
        return raw_json_text
    except Exception as e:
        logging.error(f"Erreur lors de l'appel à l'API OpenAI : {e}")
        return None

def clean_and_save_json(raw_json_text, file_path):
    """
    Nettoie le JSON renvoyé par l'API (en le parsant) et le sauvegarde dans un fichier.
    """
    try:
        clean_json_data = json.loads(raw_json_text)
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(clean_json_data, f, ensure_ascii=False, indent=4)
        logging.info(f"Fichier JSON sauvegardé à {file_path}.")
    except json.JSONDecodeError as e:
        logging.error(f"Erreur de décodage JSON : {e}")

def generate_pdf_filename(json_data, original_filename):
    """
    Génère un nouveau nom de fichier PDF à partir du nom d'origine.
    """
    base_name, _ = os.path.splitext(original_filename)
    return f"{base_name}_output.pdf"

def generate_pdf_from_json(json_data, output_file):
    """
    Génère un PDF à partir des données JSON extraites.
    """
    doc = SimpleDocTemplate(output_file, pagesize=A4)
    styles = getSampleStyleSheet()
    styles['Normal'].fontName = 'Helvetica'
    styles['Normal'].fontSize = 10
    styles['Normal'].leading = 12
    styles.add(ParagraphStyle(name='Section', parent=styles['Normal'],
                              fontName='Helvetica-Bold', fontSize=12,
                              textColor=colors.black, alignment=1, spaceAfter=12))
    styles.add(ParagraphStyle(name='Bold', parent=styles['Normal'],
                              fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='Center', parent=styles['Normal'], alignment=1))
    styles.add(ParagraphStyle(name='Right', parent=styles['Normal'], alignment=2))
    
    def create_section_title(title_text):
        title_para = Paragraph(title_text, styles['Section'])
        title_table = Table([[title_para]], colWidths=[6*inch])
        title_table.setStyle(TableStyle([
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("BOX", (0,0), (-1,-1), 1, colors.turquoise),
            ("BOTTOMPADDING", (0,0), (-1,-1), 12),
            ("TOPPADDING", (0,0), (-1,-1), 12)
        ]))
        return title_table

    story = []
    
    def draw_banner(canvas_obj, doc_obj):
        """
        Dessine la bannière en haut de la première page.
        """
        try:
            banner_path = "Background.png"  # Mettez ici le chemin vers votre bannière, si nécessaire
            if os.path.exists(banner_path):
                banner_img = ImageReader(banner_path)
                banner_height = 2 * inch
                canvas_obj.drawImage(banner_img, 0, A4[1] - banner_height,
                                     width=A4[0], height=banner_height)
            else:
                logging.warning(f"Bannière introuvable : {banner_path}")
                banner_height = 2 * inch
        except Exception as e:
            logging.error(f"Erreur lors du chargement de la bannière : {e}")
            banner_height = 2 * inch

        title_style = ParagraphStyle(
            name='HeaderTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=14,
            alignment=0,
            leading=24,
            textColor=colors.white
        )
        
        job_title = json_data.get('job_title', '').replace('\n', ' ').strip() or "CV Title"
        title_para = Paragraph(job_title, title_style)
        available_width = A4[0] - 2 * inch
        w, h = title_para.wrap(available_width, 100)
        vertical_offset = 20
        title_y = A4[1] - (banner_height)/2 - h/2 + vertical_offset
        header_x = 4.5 * cm
        title_para.drawOn(canvas_obj, header_x, title_y)
        
        full_name = json_data.get('full_name', '').strip()
        if full_name:
            parts = full_name.split()
            if len(parts) >= 2:
                initials = parts[0][0].upper() + "." + parts[-1][0].upper()
            else:
                initials = full_name[0].upper()
        else:
            initials = "?"
        canvas_obj.setFont("Helvetica-Bold", 14)
        years_experience = str(json_data.get('years_of_experience', '')).replace('\n', ' ').strip()
        experience_text = f"{initials} : {years_experience} XP"
        canvas_obj.setFillColor(colors.white)
        canvas_obj.drawString(header_x, title_y - 20, experience_text)
        
        # Exemple d'infos sur le header
        icon_y_position = A4[1] - inch - 60
        canvas_obj.setFont("Helvetica", 8)
        canvas_obj.drawString(80, icon_y_position, "01 40 76 01 49")
        canvas_obj.drawString(260, icon_y_position, "heptasys@heptasys.com")
        canvas_obj.drawString(470, icon_y_position, "www.heptasys.com")
    
    # Décalage avant le contenu principal
    story.append(Spacer(1, 100))
    
    # Formation & Certifications
    story.append(create_section_title("Formation & Certifications"))
    story.append(Spacer(1, 12))
    if json_data.get('education'):
        education_rows = []
        for edu in json_data.get('education', []):
            degree = edu.get('degree', '').replace('\n', ' ')
            institution = edu.get('institution', '').replace('\n', ' ')
            year = str(edu.get('year_of_completion', '')).strip()
            left_text = Paragraph(f"<b>{degree} à {institution}</b>", styles['Normal'])
            right_text = Paragraph(year, styles['Normal'])
            education_rows.append([left_text, right_text])
        if education_rows:
            edu_table = Table(education_rows, colWidths=[4.5*inch, 1.5*inch])
            edu_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6)
            ]))
            story.append(edu_table)
            story.append(Spacer(1, 12))
    if json_data.get('certifications'):
        certification_rows = [[Paragraph(cert, styles['Bold'])] for cert in json_data.get('certifications', [])]
        if certification_rows:
            certs_table = Table(certification_rows, colWidths=[6*inch])
            certs_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6)
            ]))
            story.append(certs_table)
            story.append(Spacer(1, 12))
    
    # Compétences techniques
    story.append(create_section_title("Compétences techniques"))
    story.append(Spacer(1, 12))
    skills_section = json_data.get('skills', None)
    if isinstance(skills_section, dict) and skills_section:
        for category_key, skills in skills_section.items():
            cat_title = category_key.replace('_', ' ').title()
            if isinstance(skills, str):
                skills = skills.strip()
            if skills:
                text = f"<b>{cat_title} :</b> {skills}"
            else:
                text = f"<b>{cat_title}</b>"
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 6))
    elif isinstance(skills_section, str) and skills_section:
        text = skills_section.strip()
        if text.endswith(":") and len(text.split(":")[-1].strip()) == 0:
            text = text[:-1].strip()
        story.append(Paragraph(text, styles['Normal']))
    else:
        story.append(Paragraph("Aucune compétence technique extraite.", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Expériences professionnelles
    story.append(create_section_title("Expériences professionnelles"))
    story.append(Spacer(1, 12))
    for exp in json_data.get('professional_experience', []):
        date_range = exp.get('date_range', '').strip() or "Date non renseignée"
        company_name = exp.get('company_name', '').replace('\n', ' ').strip()
        mission = exp.get('mission', '').replace('\n', ' ').strip() or "Poste non renseigné"
        
        exp_table = Table([
            [Paragraph(f"<b>{date_range}</b>", styles['Normal']),
             Paragraph(f"<b>{company_name}</b>", styles['Normal'])]
        ], colWidths=[3.0 * inch, 3.0 * inch])
        exp_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6)
        ]))
        story.append(exp_table)
        
        story.append(Paragraph(mission, styles['Bold']))
        
        if exp.get('tasks'):
            story.append(Paragraph("Tâches :", styles['Bold']))
            for task in exp.get('tasks'):
                story.append(Paragraph("• " + task.replace('\n', ' '), styles['Normal']))
        
        if exp.get('tech_tools'):
            tech_tools = ", ".join(exp.get('tech_tools', []))
            story.append(Paragraph("<font color='turquoise'><b>Outils</b></font> : " + tech_tools, styles['Normal']))
        
        story.append(Spacer(1, 12))
    
    doc.build(story, onFirstPage=draw_banner)

def remove_blank_paragraphs(docx_file_path):
    """
    Supprime les paragraphes vides dans un DOCX, sauf ceux contenant des images (w:drawing).
    """
    doc = Document(docx_file_path)
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip():
            if paragraph._element.xpath('.//w:drawing'):
                continue
            p = paragraph._element
            p.getparent().remove(p)
    doc.save(docx_file_path)

def convert_pdf_to_docx(pdf_path, docx_path):
    """
    Convertit un PDF en DOCX via pdf2docx, puis supprime les paragraphes vides.
    """
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0)
        cv.close()
        remove_blank_paragraphs(docx_file_path=docx_path)
        logging.info(f"Conversion réussie : {pdf_path} -> {docx_path}")
        return True
    except Exception as e:
        logging.error(f"Erreur lors de la conversion du PDF en DOCX : {e}")
        return False

def adjust_docx_top_margin(docx_file_path, top_margin_inch=0.5):
    """
    Ajuste la marge supérieure d'un DOCX.
    """
    doc = Document(docx_file_path)
    for section in doc.sections:
        section.top_margin = Inches(top_margin_inch)
    doc.save(docx_file_path)

def upload_to_blob_storage(file_path, blob_name):
    """
    Upload un fichier local vers Azure Blob Storage dans le conteneur défini.
    """
    try:
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)
        with open(file_path, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
        logging.info(f"Fichier {blob_name} uploadé vers Azure Blob Storage.")
    except Exception as e:
        logging.error(f"Erreur lors de l'upload vers Blob Storage : {e}")

def generate_sas_token(blob_name):
    """
    Génère un SAS token en lecture seule (valable 10 minutes) pour un blob.
    """
    try:
        sas_token = generate_blob_sas(
            account_name=blob_service_client.account_name,
            container_name=container_name,
            blob_name=blob_name,
            account_key=_account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.now(timezone.utc) + timedelta(minutes=10)
        )
        blob_url = f"https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob_name}?{sas_token}"
        return blob_url
    except Exception as e:
        logging.error(f"Erreur lors de la génération du SAS token : {e}")
        return None

@app.route('/generate-sas-token', methods=['POST'])
def generate_sas_token_route():
    """
    Endpoint pour générer un SAS token en fournissant le nom du blob via JSON.
    """
    data = request.json
    blob_name = data.get('blob_name')
    if not blob_name:
        return jsonify({"error": "Le nom du blob est requis"}), 400
    sas_url = generate_sas_token(blob_name)
    if sas_url:
        return jsonify({"sas_url": sas_url}), 200
    else:
        return jsonify({"error": "Échec de la génération du SAS token"}), 500

def allowed_file(filename):
    """
    Vérifie l'extension du fichier : PDF, DOCX, PNG, JPG, JPEG.
    """
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'png', 'jpg', 'jpeg'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/template', methods=['POST'])
def upload_file():
    """
    Endpoint principal : 
    1) Récupère le fichier depuis la requête POST,
    2) Extrait le texte et envoie au modèle AzureOpenAI,
    3) Convertit la réponse en JSON, 
    4) Génère un PDF et un DOCX,
    5) Upload les deux sur Blob Storage, 
    6) Retourne les URLs SAS en JSON.
    """
    logging.info("Requête reçue sur /template")
    
    # Vérification de la présence du fichier
    if 'file' not in request.files:
        logging.error("Aucun fichier trouvé dans la requête")
        return jsonify({"error": "Aucun fichier trouvé dans la requête"}), 400
    
    file = request.files['file']
    if file.filename == '':
        logging.error("Aucun fichier sélectionné")
        return jsonify({"error": "Aucun fichier sélectionné"}), 400
    
    # Vérification de l’extension
    if file and allowed_file(file.filename):
        filename = file.filename
        file_path = os.path.join(tempfile.gettempdir(), filename)
        file.save(file_path)
        logging.info(f"Fichier sauvegardé à {file_path}")
        
        # Extraction de texte
        extracted_text = extract_text(file_path)
        if not extracted_text:
            logging.error("Aucun texte extrait du fichier")
            return jsonify({"error": "Échec de l'extraction du texte"}), 500
        logging.info("Texte extrait du fichier avec succès.")
        
        # Extraction des informations (JSON)
        raw_json_text = extract_info_to_json(extracted_text)
        if not raw_json_text:
            logging.error("Échec de l'extraction des informations structurées (réponse vide)")
            return jsonify({"error": "Échec de l'extraction des informations structurées"}), 500
        
        logging.info("Informations extraites au format JSON (brut).")
        
        # Nettoyage & sauvegarde JSON
        json_file_path = os.path.join(tempfile.gettempdir(), 'extracted_info.json')
        clean_and_save_json(raw_json_text, json_file_path)
        
        if not os.path.exists(json_file_path):
            logging.error(f"Fichier JSON non trouvé à {json_file_path}")
            return jsonify({"error": "Échec de la sauvegarde du fichier JSON"}), 500
        
        with open(json_file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        logging.info("Données JSON chargées : %s", json_data)
        
        # Génération du PDF
        pdf_file_name = generate_pdf_filename(json_data, filename)
        pdf_file_path = os.path.join(tempfile.gettempdir(), pdf_file_name)
        generate_pdf_from_json(json_data, pdf_file_path)
        logging.info(f"PDF généré à {pdf_file_path}")
        
        # Conversion PDF → DOCX
        docx_file_name = pdf_file_name.replace('.pdf', '.docx')
        docx_file_path = os.path.join(tempfile.gettempdir(), docx_file_name)
        if not convert_pdf_to_docx(pdf_file_path, docx_file_path):
            logging.error("Échec de la conversion du PDF en DOCX")
            return jsonify({"error": "Échec de la conversion du PDF en DOCX"}), 500
        
        logging.info(f"DOCX généré à {docx_file_path}")
        
        # Ajustement de la marge supérieure
        adjust_docx_top_margin(docx_file_path, top_margin_inch=0.5)
        logging.info("Marge supérieure et espacement ajustés dans le fichier DOCX.")
        
        # Upload dans le Blob Storage
        upload_to_blob_storage(pdf_file_path, pdf_file_name)
        upload_to_blob_storage(docx_file_path, docx_file_name)
        logging.info(f"PDF et DOCX uploadés vers Blob Storage sous les noms {pdf_file_name} et {docx_file_name}")
        
        # Génération des URLs SAS
        pdf_sas_url = generate_sas_token(pdf_file_name)
        docx_sas_url = generate_sas_token(docx_file_name)
        
        if pdf_sas_url and docx_sas_url:
            logging.info(f"SAS URLs générés : PDF - {pdf_sas_url}, DOCX - {docx_sas_url}")
            return jsonify({
                "pdf_sas_url": pdf_sas_url,
                "docx_sas_url": docx_sas_url
            }), 200
        else:
            logging.error("Échec de la génération des SAS tokens")
            return jsonify({"error": "Échec de la génération des SAS tokens"}), 500
    
    else:
        logging.error("Fichier non valide ou extension non autorisée")
        return jsonify({"error": "Fichier non valide ou extension non autorisée"}), 400

if __name__ == "__main__":
    # Démarre l’application Flask sur le port 5001
    app.run(host='0.0.0.0', port=5001, debug=True)
